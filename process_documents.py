import os
import io
import docx
import json
import time
from dotenv import load_dotenv
from openai import OpenAI # Nova importação para a OpenAI

# Carrega as variáveis de ambiente do arquivo .env
load_dotenv()

# Configura a API da OpenAI com a chave
# O cliente OpenAI lê a variável de ambiente OPENAI_API_KEY por padrão
client = OpenAI() 
if not os.getenv("OPENAI_API_KEY"):
    raise ValueError("A variável de ambiente OPENAI_API_KEY não foi encontrada.")

# Constantes do projeto (permanecem as mesmas)
# ... (imports do Google, constantes de pastas, etc)
from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
from googleapiclient.http import MediaIoBaseDownload

SCOPES = ["https://www.googleapis.com/auth/drive.readonly"]
FOLDER_ID_PAGOS = "1jTRfpGeotGcd3YZZA-4YvwAxIrdGp14G"
FOLDER_ID_GRATUITOS = "1NBwjHCLjpIIh04p7sKGBqHCODeXIEg7p"


def get_drive_service():
    """Autentica e retorna um objeto de serviço da API do Drive."""
    creds = None
    if os.path.exists("token.json"):
        creds = Credentials.from_authorized_user_file("token.json", SCOPES)
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            raise Exception("Token não encontrado ou inválido. Execute 'authenticate.py' primeiro.")
        with open("token.json", "w") as token:
            token.write(creds.to_json())
    return build("drive", "v3", credentials=creds)

def get_docx_text(service, file_id):
    """Faz o download de um arquivo .docx e extrai todo o seu texto."""
    try:
        request = service.files().get_media(fileId=file_id)
        file_buffer = io.BytesIO()
        downloader = MediaIoBaseDownload(file_buffer, request)
        done = False
        while not done:
            status, done = downloader.next_chunk()
        file_buffer.seek(0)
        document = docx.Document(file_buffer)
        full_text = []
        for para in document.paragraphs:
            full_text.append(para.text)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        return '\n'.join(full_text)
    except HttpError as error:
        print(f"    ERRO ao baixar o arquivo: {error}")
        return None

def extract_data_with_openai(text, file_name):
    """
    Usa a API da OpenAI (GPT-4o) para extrair dados estruturados.
    """
    print("    - Enviando para a API da OpenAI (GPT-4o)...")
    
    system_prompt = """
    Você é um assistente de IA altamente preciso, especializado em extrair dados de documentos contratuais e formatá-los como um objeto JSON.
    Sua resposta deve conter APENAS o objeto JSON, sem nenhum texto, explicação ou ```json``` adicional.
    Se um campo não for encontrado no texto, seu valor deve ser `null`.
    """

    user_prompt = f"""
    Analise o seguinte "Termo de Permissão de Uso" (nome do arquivo: {file_name}) e extraia as informações conforme a estrutura JSON solicitada.

    Instruções Específicas:
    - `cliente.documento`: Retorne apenas os números do CNPJ ou CPF.
    - `eventos.valor_final`: Retorne um número (float). Se o evento for gratuito, retorne 0.0.
    - `eventos.datas_evento`: Retorne uma lista de strings, com cada data no formato "YYYY-MM-DD".

    Texto do Documento:
    ---
    {text}
    ---

    Estrutura JSON de Saída:
    {{
      "cliente": {{
        "nome_razao_social": "string",
        "documento": "string",
        "tipo_pessoa": "string ('PJ' ou 'PF')",
        "nome_responsavel": "string"
      }},
      "eventos": [
        {{
          "numero_processo": "string",
          "numero_termo": "string",
          "nome_evento": "string",
          "datas_evento": ["YYYY-MM-DD"],
          "hora_inicio": "string",
          "hora_fim": "string",
          "valor_final": 0.0,
          "espaco_utilizado": "string"
        }}
      ]
    }}
    """
    
    max_retries = 3
    retry_delay = 5
    for attempt in range(max_retries):
        try:
            response = client.chat.completions.create(
                model="gpt-4o",  # Ou "gpt-4-turbo"
                response_format={"type": "json_object"},
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ]
            )
            print("    - Resposta recebida da IA.")
            return json.loads(response.choices[0].message.content)
        except Exception as e:
            print(f"    ERRO ao chamar a API da OpenAI (tentativa {attempt + 1}/{max_retries}): {e}")
            if attempt < max_retries - 1:
                print(f"    - Esperando {retry_delay} segundos para tentar novamente...")
                time.sleep(retry_delay)
                retry_delay *= 2
            else:
                print("    - Falha ao extrair dados após múltiplas tentativas.")
                return None

def main():
    """
    Script principal para extrair dados de todos os documentos usando a API da OpenAI.
    """
    print("Iniciando o processo de extração com IA (Modelo: OpenAI GPT-4o)...")
    service = get_drive_service()
    
    folder_ids = {
        "Termos Pagos": FOLDER_ID_PAGOS,
        "Termos Gratuitos": FOLDER_ID_GRATUITOS
    }
    all_extracted_data = []

    for folder_name, folder_id in folder_ids.items():
        print(f"\n--- Processando pasta: {folder_name} ---")
        try:
            query = (f"'{folder_id}' in parents and "
                     "mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document' and "
                     "trashed=false")
            results = service.files().list(q=query, pageSize=1000, fields="files(id, name)").execute()
            items = results.get("files", [])

            if not items:
                print(f"Nenhum arquivo .docx encontrado na pasta {folder_name}.")
                continue

            print(f"Encontrados {len(items)} arquivos. Iniciando extração com OpenAI...")
            for i, item in enumerate(items):
                print(f"\n[ Processando {i+1}/{len(items)} ] Lendo arquivo: {item['name']}")
                text_content = get_docx_text(service, item['id'])
                
                if text_content:
                    extracted_data = extract_data_with_openai(text_content, item['name'])
                    if extracted_data:
                        extracted_data['arquivo_origem'] = item['name']
                        extracted_data['id_arquivo_drive'] = item['id']
                        all_extracted_data.append(extracted_data)
                        print("    - Extração bem-sucedida.")
                    else:
                        print("    - Falha na extração com IA.")
                
                # Pausa para não sobrecarregar a API
                time.sleep(1) 

        except HttpError as error:
            print(f"Ocorreu um erro ao acessar a pasta {folder_name}: {error}")
    
    output_filename = 'dados_extraidos_openai.json'
    with open(output_filename, 'w', encoding='utf-8') as f:
        json.dump(all_extracted_data, f, indent=2, ensure_ascii=False)
    
    print(f"\n\nProcesso concluído! Todos os dados extraídos pela OpenAI foram salvos em '{output_filename}'.")

if __name__ == "__main__":
    main()